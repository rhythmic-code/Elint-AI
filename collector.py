# import imaplib
# import email
# import os
# import re
# import hashlib
# import logging
# from email.header import decode_header
# from PyPDF2 import PdfReader
# from docx import Document
# from dotenv import load_dotenv
# from email.utils import parsedate_to_datetime
# from datetime import datetime

# # Load environment variables
# load_dotenv()

# # Configure logging (both file and console)
# logging.basicConfig(
#     filename='resume_processor.log',
#     level=logging.INFO,
#     format='%(asctime)s - %(levelname)s - %(message)s'
# )
# console_handler = logging.StreamHandler()
# console_handler.setLevel(logging.INFO)
# formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
# console_handler.setFormatter(formatter)
# logging.getLogger().addHandler(console_handler)

# # Directory to save resumes
# RESUME_DIR = 'resumes'
# os.makedirs(RESUME_DIR, exist_ok=True)

# # Regular expressions for masking PII
# EMAIL_REGEX = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
# PHONE_REGEX = re.compile(r'(\+?\d{1,2}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}')
# NAME_REGEX = re.compile(r'\b([A-Z][a-z]{2,}\s[A-Z][a-z]{2,}(?:\s[A-Z][a-z]{2,})?)\b')

# # Patterns to extract batch year
# BATCH_PATTERNS = [
#     r'Batch\s+of\s+(\d{4})',
#     r'Graduated\s+(\d{4})',
#     r'Class\s+of\s+(\d{4})',
#     r'(\d{4})\s+Batch'
# ]

# # AI-related keywords
# AI_KEYWORDS = [
#     'machine learning', 'deep learning', 'artificial intelligence',
#     'neural networks', 'nlp', 'computer vision', 'tensorflow',
#     'pytorch', 'scikit-learn', 'data science', 'ml', 'ai'
# ]

# def mask_pii(text):
#     """Mask personal identifiable information in the text."""
#     text = EMAIL_REGEX.sub('[EMAIL]', text)
#     text = PHONE_REGEX.sub('[PHONE]', text)
#     text = NAME_REGEX.sub('[NAME]', text)
#     return text

# def extract_batch_year(text):
#     """Extract batch year from resume text."""
#     for pattern in BATCH_PATTERNS:
#         match = re.search(pattern, text, re.IGNORECASE)
#         if match:
#             return match.group(1)
#     return "Not found"

# def extract_ai_experience(text):
#     """Extract AI-related experience from resume text."""
#     experience = []
#     lines = text.split('\n')
#     for line in lines:
#         if any(keyword in line.lower() for keyword in AI_KEYWORDS):
#             experience.append(line.strip())
#     return experience

# def extract_text(file_path):
#     """Extract text from PDF or DOCX files."""
#     try:
#         if file_path.endswith('.pdf'):
#             with open(file_path, 'rb') as f:
#                 reader = PdfReader(f)
#                 return "\n".join(page.extract_text() or "" for page in reader.pages)
#         elif file_path.endswith('.docx'):
#             doc = Document(file_path)
#             return '\n'.join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         logging.error(f"Error extracting text from {file_path}: {str(e)}")
#     return ""

# def fetch_resumes_with_senders(limit=50):
#     server = os.getenv("IMAP_SERVER", "imap.gmail.com")
#     username = os.getenv("EMAIL_USER")
#     password = os.getenv("EMAIL_PASS")

#     if not username or not password:
#         logging.error("Email credentials not set in environment.")
#         return []

#     try:
#         with imaplib.IMAP4_SSL(server) as mail:
#             mail.login(username, password)
#             mail.select('inbox')
#             _, data = mail.search(None, 'UNSEEN')  # Fetch only unread emails
#             email_ids = data[0].split()
#             email_ids = email_ids[-limit:]

#             resumes = []

#             for num in email_ids:
#                 _, msg_data = mail.fetch(num, '(RFC822)')
#                 msg = email.message_from_bytes(msg_data[0][1])
#                 sender_email = email.utils.parseaddr(msg["From"])[1]

#                 try:
#                     date_str = parsedate_to_datetime(msg['Date']).strftime("%Y%m%d_%H%M%S")
#                 except Exception:
#                     date_str = datetime.now().strftime("%Y%m%d_%H%M%S")

#                 for part in msg.walk():
#                     if part.get_content_disposition() == 'attachment':
#                         filename = part.get_filename()
#                         if filename:
#                             try:
#                                 decoded = decode_header(filename)[0]
#                                 filename = decoded[0].decode(decoded[1]) if isinstance(decoded[0], bytes) else decoded[0]
#                             except Exception as e:
#                                 logging.warning(f"Failed to decode filename: {str(e)}")
#                                 continue

#                             if not filename.lower().endswith((".pdf", ".docx")):
#                                 continue

#                             unique_hash = hashlib.md5((sender_email + date_str + filename).encode()).hexdigest()
#                             save_name = f"{sender_email}_{date_str}_{unique_hash}_{filename}"
#                             file_path = os.path.join(RESUME_DIR, save_name)

#                             with open(file_path, 'wb') as f:
#                                 f.write(part.get_payload(decode=True))

#                             resumes.append({
#                                 "path": file_path,
#                                 "sender": sender_email,
#                                 "timestamp": date_str
#                             })

#                 mail.store(num, '+FLAGS', '\\Seen')  # Mark as read

#             return resumes

#     except Exception as e:
#         logging.error(f"Error fetching emails: {str(e)}")
#         return []

# def process_resumes(resume_records):
#     """Process each resume to extract and mask information."""
#     for record in resume_records:
#         path = record["path"]
#         text = extract_text(path)
#         if not text:
#             continue
#         masked_text = mask_pii(text)
#         batch_year = extract_batch_year(text)
#         ai_experience = extract_ai_experience(text)

#         print(f"Processed Resume: {os.path.basename(path)}")
#         print(f"Masked Text:\n{masked_text[:500]}...")  # Display first 500 characters
#         print(f"Batch Year: {batch_year}")
#         print(f"AI Experience: {ai_experience}")
#         print("-" * 50)

# def run_collector():
#     resumes = fetch_resumes_with_senders(limit=50)
#     if resumes:
#         process_resumes(resumes)
#         logging.info(f"Total resumes processed: {len(resumes)}")
#     else:
#         logging.info("No new resumes found.")

# if __name__ == "__main__":
#     run_collector()
# collector.py
import imaplib
import email
import os
import re
import hashlib
import logging
from email.header import decode_header
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
from email.utils import parsedate_to_datetime
from datetime import datetime

load_dotenv()

logging.basicConfig(
    filename='resume_processor.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console_handler.setFormatter(formatter)
logging.getLogger().addHandler(console_handler)

RESUME_DIR = 'resumes'
os.makedirs(RESUME_DIR, exist_ok=True)

EMAIL_REGEX = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
PHONE_REGEX = re.compile(r'(\+?\d{1,2}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}')
NAME_REGEX  = re.compile(r'\b([A-Z][a-z]{2,}\s[A-Z][a-z]{2,}(?:\s[A-Z][a-z]{2,})?)\b')

BATCH_PATTERNS = [
    r'Batch\s+of\s+(\d{4})',
    r'Graduated\s+(\d{4})',
    r'Class\s+of\s+(\d{4})',
    r'(\d{4})\s+Batch'
]

AI_KEYWORDS = [
    'machine learning', 'deep learning', 'artificial intelligence',
    'neural networks', 'nlp', 'computer vision', 'tensorflow',
    'pytorch', 'scikit-learn', 'data science', 'ml', 'ai'
]

def mask_pii(text):
    text = EMAIL_REGEX.sub('[EMAIL]', text)
    text = PHONE_REGEX.sub('[PHONE]', text)
    text = NAME_REGEX.sub('[NAME]',  text)
    return text

def extract_batch_year(text):
    for pattern in BATCH_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
    return "Not found"

def extract_ai_experience(text):
    exp = []
    for line in text.splitlines():
        if any(kw in line.lower() for kw in AI_KEYWORDS):
            exp.append(line.strip())
    return exp

def extract_text(path):
    try:
        if path.lower().endswith('.pdf'):
            with open(path, 'rb') as f:
                reader = PdfReader(f)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif path.lower().endswith('.docx'):
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logging.error(f"Error extracting text from {path}: {e}")
    return ""

def fetch_resumes(limit=50):
    """Log in, pull up to `limit` new resume attachments, save them locally."""
    server   = os.getenv("IMAP_SERVER", "imap.gmail.com")
    user     = os.getenv("EMAIL_USER")
    password = os.getenv("EMAIL_PASS")

    if not user or not password:
        logging.error("Email credentials missing.")
        return []

    try:
        with imaplib.IMAP4_SSL(server) as mail:
            mail.login(user, password)
            mail.select('inbox')
            _, data = mail.search(None, 'UNSEEN')
            ids = data[0].split()[-limit:]
            records = []

            for num in ids:
                _, raw = mail.fetch(num, '(RFC822)')
                msg = email.message_from_bytes(raw[0][1])
                sender = email.utils.parseaddr(msg["From"])[1]
                try:
                    dt = parsedate_to_datetime(msg['Date']).strftime("%Y%m%d_%H%M%S")
                except:
                    dt = datetime.now().strftime("%Y%m%d_%H%M%S")

                for part in msg.walk():
                    if part.get_content_disposition() == 'attachment':
                        fn = part.get_filename()
                        if not fn or not fn.lower().endswith(('.pdf','.docx', '.jpg', '.png', 'jpeg')):
                            continue
                        decoded = decode_header(fn)[0]
                        fn = decoded[0].decode(decoded[1]) if isinstance(decoded[0], bytes) else decoded[0]

                        h = hashlib.md5((sender+dt+fn).encode()).hexdigest()
                        save_as = f"{sender}_{dt}_{h}_{fn}"
                        path = os.path.join(RESUME_DIR, save_as)

                        with open(path, 'wb') as out:
                            out.write(part.get_payload(decode=True))
                        records.append({'path': path, 'sender': sender, 'timestamp': dt})

                mail.store(num, '+FLAGS', '\\Seen')
            return records

    except Exception as e:
        logging.error(f"Error fetching resumes: {e}")
        return []

def process_resumes(records):
    """Mask PII, extract batch & AI exp, and log summary to console."""
    for rec in records:
        text = extract_text(rec['path'])
        if not text:
            continue
        masked = mask_pii(text)
        rec['batch_year']    = extract_batch_year(text)
        rec['ai_experience'] = extract_ai_experience(text)
        # just a quick console summary:
        print(f"→ {os.path.basename(rec['path'])}")
        print(f"   Batch: {rec['batch_year']}")
        print(f"   AI exp: {len(rec['ai_experience'])} lines")
        print("-"*40)

# Note: no if __name__ guard here
